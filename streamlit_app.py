import streamlit as st
import pandas as pd
import numpy as np
from docxtpl import DocxTemplate
from datetime import datetime
import io
import zipfile
import re
import docx
import json
import openpyxl
from openpyxl.utils import get_column_letter
import tempfile
import base64
import plotly.express as px
import plotly.graph_objects as go
from typing import Dict, List, Any, Optional

# --- Page Configuration ---
st.set_page_config(
    page_title="Advanced Document Assembly System",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Advanced Classes from newlettergen.py ---

class ExcelFormatHandler:
    """Handles Excel format detection and application for better data presentation"""
    
    def __init__(self):
        self.format_patterns = {
            '[$£-809]': {'type': 'currency', 'symbol': '£', 'thousands_sep': True},
            '[$€-2]': {'type': 'currency', 'symbol': '€', 'thousands_sep': True},
            r'[$\$-409]': {'type': 'currency', 'symbol': '$', 'thousands_sep': True},
            '£': {'type': 'currency', 'symbol': '£', 'thousands_sep': True},
            '$': {'type': 'currency', 'symbol': '$', 'thousands_sep': True},
            '€': {'type': 'currency', 'symbol': '€', 'thousands_sep': True},
            '0%': {'type': 'percentage', 'decimals': 0},
            '0.00%': {'type': 'percentage', 'decimals': 2},
            '%': {'type': 'percentage', 'decimals': 2},
            '#,##0': {'type': 'number', 'thousands_sep': True, 'decimals': 0},
            '#,##0.00': {'type': 'number', 'thousands_sep': True, 'decimals': 2},
        }
    
    def format_value(self, value, excel_format=None, format_type=None, column_name=None):
        """Format a value based on Excel format or detected format type"""
        if pd.isna(value):
            return ""
        
        # Auto-detect format from column name or value
        if format_type == "Auto" or not format_type:
            if column_name:
                if any(term in column_name.lower() for term in ['price', 'cost', 'value', 'amount']):
                    format_type = "Currency (£)"
                elif '%' in column_name or 'percent' in column_name.lower():
                    format_type = "Percentage (%)"
        
        # Apply formatting
        if format_type == "Currency (£)":
            try:
                num_val = float(value)
                return f"£{num_val:,.2f}"
            except:
                return str(value)
        elif format_type == "Percentage (%)":
            try:
                num_val = float(value)
                if 0 <= num_val <= 1:
                    num_val *= 100
                return f"{num_val:.1f}%"
            except:
                return str(value)
        elif format_type == "Number":
            try:
                num_val = float(value)
                if num_val == int(num_val):
                    return f"{int(num_val):,}"
                return f"{num_val:,.2f}"
            except:
                return str(value)
        elif format_type == "Date":
            try:
                if isinstance(value, pd.Timestamp):
                    return value.strftime("%d/%m/%Y")
                return pd.to_datetime(value).strftime("%d/%m/%Y")
            except:
                return str(value)
        
        return str(value)

# --- Initialize Session State ---
def initialize_session_state():
    """Initialize all session state variables"""
    defaults = {
        'df': None,
        'current_sheet': None,
        'excel_sheets': [],
        'detected_fields': [],
        'field_mappings': {},
        'conditional_rules': {},
        'client_filters': [],
        'table_definitions': [],
        'cell_formats': {},
        'uploaded_data_name': None,
        'uploaded_template_name': None,
        'format_handler': ExcelFormatHandler(),
        'preview_client': None,
        'generation_progress': 0,
        'generation_status': "Ready",
        'template_analysis_debug': [],
        'on_the_fly_placeholders': {}  # Fixed: Added missing initialization
    }
    
    for key, default_value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

# --- Helper Functions ---

def load_excel_with_formats(file_content, sheet_name=None):
    """Load Excel data and extract formatting information"""
    try:
        # Load data
        df = pd.read_excel(io.BytesIO(file_content), sheet_name=sheet_name)
        
        # Extract format information using openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_content), data_only=False)
        ws = wb[sheet_name] if sheet_name else wb.active
        
        cell_formats = {}
        for row_idx in range(2, ws.max_row + 1):  # Skip header
            for col_idx in range(1, ws.max_column + 1):
                cell = ws.cell(row_idx, col_idx)
                if cell.number_format and cell.number_format != 'General':
                    col_letter = get_column_letter(col_idx)
                    col_name = ws.cell(1, col_idx).value  # Get header name
                    if col_name in df.columns:
                        df_row_idx = row_idx - 2  # Adjust for pandas indexing
                        cell_formats[(df_row_idx, col_name)] = cell.number_format
        
        return df, cell_formats, wb.sheetnames
    except Exception as e:
        st.error(f"Error loading Excel file: {e}")
        return None, {}, []

def detect_template_fields(template_file):
    """Detect fields in Word template with enhanced debugging"""
    if not template_file:
        st.session_state.template_analysis_debug = ["No template file provided"]
        return []
    
    debug_info = []
    all_fields = []
    
    try:
        # Reset file position
        template_file.seek(0)
        file_content = template_file.read()
        debug_info.append(f"Template file size: {len(file_content)} bytes")
        
        # Try Jinja placeholders first
        template_buffer = io.BytesIO(file_content)
        
        try:
            doc = DocxTemplate(template_buffer)
            jinja_fields = list(doc.get_undeclared_template_variables())
            debug_info.append(f"Jinja fields found: {jinja_fields}")
            all_fields.extend(jinja_fields)
        except Exception as e:
            debug_info.append(f"Jinja detection failed: {str(e)}")
            jinja_fields = []
        
        # Detect various bracket-style placeholders
        template_buffer = io.BytesIO(file_content)
        
        try:
            docx_doc = docx.Document(template_buffer)
            bracket_fields = []
            
            # Enhanced patterns to match multiple placeholder styles
            patterns = {
                'single_curly': r'\{([^{}]+)\}',           # {field}
                'double_curly': r'\{\{([^{}]+)\}\}',       # {{field}}
                'triple_curly': r'\{\{\{([^{}]+)\}\}\}',   # {{{field}}}
                'single_square': r'\[([^\[\]]+)\]',        # [field]
                'double_square': r'\[\[([^\[\]]+)\]\]',    # [[field]]
                'single_angle': r'<([^<>]+)>',             # <field>
                'double_angle': r'<<([^<>]+)>>',           # <<field>>
                'triple_angle': r'<<<([^<>]+)>>>',         # <<<field>>>
                'double_paren': r'\(\(([^()]+)\)\)',       # ((field))
                'triple_paren': r'\(\(\(([^()]+)\)\)\)',   # (((field)))
            }
            
            total_text = ""
            
            # Search in paragraphs
            for paragraph in docx_doc.paragraphs:
                total_text += paragraph.text + " "
                for pattern_name, pattern in patterns.items():
                    matches = re.findall(pattern, paragraph.text)
                    for match in matches:
                        field = match.strip() if isinstance(match, str) else match
                        if field and field not in jinja_fields and field not in bracket_fields:
                            bracket_fields.append(field)
                            debug_info.append(f"Found {pattern_name} field: '{field}' in paragraph")
            
            # Search in table cells
            for table_idx, table in enumerate(docx_doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        total_text += cell.text + " "
                        for pattern_name, pattern in patterns.items():
                            matches = re.findall(pattern, cell.text)
                            for match in matches:
                                field = match.strip() if isinstance(match, str) else match
                                if field and field not in jinja_fields and field not in bracket_fields:
                                    bracket_fields.append(field)
                                    debug_info.append(f"Found {pattern_name} field: '{field}' in table {table_idx}, row {row_idx}, cell {cell_idx}")
            
            debug_info.append(f"Total document text length: {len(total_text)} characters")
            debug_info.append(f"Bracket fields found: {bracket_fields}")
            all_fields.extend(bracket_fields)
            
        except Exception as e:
            debug_info.append(f"Bracket detection failed: {str(e)}")
        
        # Combine and deduplicate
        all_fields = list(set(all_fields))
        all_fields.sort()
        
        debug_info.append(f"Final combined fields: {all_fields}")
        
    except Exception as e:
        debug_info.append(f"Overall detection failed: {str(e)}")
        
    # Store debug info in session state
    st.session_state.template_analysis_debug = debug_info
    
    return all_fields

def convert_placeholders_to_jinja(template_buffer, detected_fields):
    """Convert all detected placeholder formats to Jinja2 format for DocxTemplate processing"""
    
    # Load the document to modify text
    temp_doc = docx.Document(template_buffer)
    
    # Define the patterns that need conversion (exclude Jinja patterns)
    conversion_patterns = {
        'single_curly': r'\{([^{}]+)\}',           # {field} -> {{field}}
        'single_square': r'\[([^\[\]]+)\]',        # [field] -> {{field}}
        'double_square': r'\[\[([^\[\]]+)\]\]',    # [[field]] -> {{field}}
        'single_angle': r'<([^<>]+)>',             # <field> -> {{field}}
        'double_angle': r'<<([^<>]+)>>',           # <<field>> -> {{field}}
        'triple_angle': r'<<<([^<>]+)>>>',         # <<<field>>> -> {{field}}
        'double_paren': r'\(\(([^()]+)\)\)',       # ((field)) -> {{field}}
        'triple_paren': r'\(\(\(([^()]+)\)\)\)',   # (((field))) -> {{field}}
        'triple_curly': r'\{\{\{([^{}]+)\}\}\}',   # {{{field}}} -> {{field}}
    }
    
    # Convert text in paragraphs
    for paragraph in temp_doc.paragraphs:
        original_text = paragraph.text
        modified_text = original_text
        
        # Apply each pattern conversion
        for pattern_name, pattern in conversion_patterns.items():
            matches = re.findall(pattern, modified_text)
            for match in matches:
                field_name = match.strip()
                # Only convert if this field is in our detected fields
                if field_name in detected_fields:
                    # Find the original placeholder text
                    if pattern_name == 'single_curly':
                        original_placeholder = f"{{{field_name}}}"
                    elif pattern_name == 'single_square':
                        original_placeholder = f"[{field_name}]"
                    elif pattern_name == 'double_square':
                        original_placeholder = f"[[{field_name}]]"
                    elif pattern_name == 'single_angle':
                        original_placeholder = f"<{field_name}>"
                    elif pattern_name == 'double_angle':
                        original_placeholder = f"<<{field_name}>>"
                    elif pattern_name == 'triple_angle':
                        original_placeholder = f"<<<{field_name}>>>"
                    elif pattern_name == 'double_paren':
                        original_placeholder = f"(({field_name}))"
                    elif pattern_name == 'triple_paren':
                        original_placeholder = f"((({field_name})))"
                    elif pattern_name == 'triple_curly':
                        original_placeholder = f"{{{{{field_name}}}}}"
                    else:
                        continue
                    
                    # Convert to Jinja2 format
                    jinja_placeholder = f"{{{{{field_name}}}}}"
                    
                    # Replace in the text
                    modified_text = modified_text.replace(original_placeholder, jinja_placeholder)
        
        # Update paragraph text if it changed
        if modified_text != original_text:
            paragraph.text = modified_text
    
    # Convert text in table cells
    for table in temp_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    modified_text = original_text
                    
                    # Apply each pattern conversion
                    for pattern_name, pattern in conversion_patterns.items():
                        matches = re.findall(pattern, modified_text)
                        for match in matches:
                            field_name = match.strip()
                            # Only convert if this field is in our detected fields
                            if field_name in detected_fields:
                                # Find the original placeholder text
                                if pattern_name == 'single_curly':
                                    original_placeholder = f"{{{field_name}}}"
                                elif pattern_name == 'single_square':
                                    original_placeholder = f"[{field_name}]"
                                elif pattern_name == 'double_square':
                                    original_placeholder = f"[[{field_name}]]"
                                elif pattern_name == 'single_angle':
                                    original_placeholder = f"<{field_name}>"
                                elif pattern_name == 'double_angle':
                                    original_placeholder = f"<<{field_name}>>"
                                elif pattern_name == 'triple_angle':
                                    original_placeholder = f"<<<{field_name}>>>"
                                elif pattern_name == 'double_paren':
                                    original_placeholder = f"(({field_name}))"
                                elif pattern_name == 'triple_paren':
                                    original_placeholder = f"((({field_name})))"
                                elif pattern_name == 'triple_curly':
                                    original_placeholder = f"{{{{{field_name}}}}}"
                                else:
                                    continue
                                
                                # Convert to Jinja2 format
                                jinja_placeholder = f"{{{{{field_name}}}}}"
                                
                                # Replace in the text
                                modified_text = modified_text.replace(original_placeholder, jinja_placeholder)
                    
                    # Update paragraph text if it changed
                    if modified_text != original_text:
                        paragraph.text = modified_text
    
    # Save modified template to buffer
    modified_template_buffer = io.BytesIO()
    temp_doc.save(modified_template_buffer)
    modified_template_buffer.seek(0)
    
    return modified_template_buffer

def apply_filter(data, column, operator, value):
    """Apply a filter to dataframe"""
    if operator == "equals":
        return data[data[column].astype(str) == value]
    elif operator == "contains":
        return data[data[column].astype(str).str.contains(value, na=False)]
    elif operator == "starts with":
        return data[data[column].astype(str).str.startswith(value, na=False)]
    elif operator == "ends with":
        return data[data[column].astype(str).str.endswith(value, na=False)]
    elif operator == "greater than":
        try:
            return data[data[column].astype(float) > float(value)]
        except:
            return data.iloc[0:0]
    elif operator == "less than":
        try:
            return data[data[column].astype(float) < float(value)]
        except:
            return data.iloc[0:0]
    elif operator == "is empty":
        return data[data[column].isna() | (data[column].astype(str) == "")]
    elif operator == "is not empty":
        return data[~data[column].isna() & (data[column].astype(str) != "")]
    return data

def evaluate_condition(cell_value, operator, value):
    """Evaluate a single condition"""
    if pd.isna(cell_value):
        cell_value = ""
    
    if operator == "equals":
        return str(cell_value).lower() == str(value).lower()
    elif operator == "does not equal":
        return str(cell_value).lower() != str(value).lower()
    elif operator == "contains":
        return str(value).lower() in str(cell_value).lower()
    elif operator == "does not contain":
        return str(value).lower() not in str(cell_value).lower()
    elif operator == "starts with":
        return str(cell_value).lower().startswith(str(value).lower())
    elif operator == "ends with":
        return str(cell_value).lower().endswith(str(value).lower())
    elif operator == "greater than":
        try:
            return float(cell_value) > float(value)
        except:
            return False
    elif operator == "less than":
        try:
            return float(cell_value) < float(value)
        except:
            return False
    elif operator == "is empty":
        return pd.isna(cell_value) or str(cell_value).strip() == ""
    elif operator == "is not empty":
        return not pd.isna(cell_value) and str(cell_value).strip() != ""
    return False

def evaluate_multiple_conditions(conditions, row_data):
    """Evaluate multiple conditions with AND/OR logic"""
    if not conditions:
        return False
    
    # First condition
    first_condition = conditions[0]
    column = first_condition.get("column", "")
    operator = first_condition.get("operator", "")
    value = first_condition.get("value", "")
    
    if column not in row_data:
        return False
    
    result = evaluate_condition(row_data[column], operator, value)
    
    # Evaluate remaining conditions
    for condition in conditions[1:]:
        logic = condition.get("logic", "AND")
        column = condition.get("column", "")
        operator = condition.get("operator", "")
        value = condition.get("value", "")
        
        if column not in row_data:
            if logic == "AND":
                return False
            continue
        
        condition_result = evaluate_condition(row_data[column], operator, value)
        
        if logic == "AND":
            result = result and condition_result
            if not result:
                return False
        else:  # OR
            result = result or condition_result
    
    return result

def generate_content_values(client_id, client_data, format_handler):
    """Generate content values for a client based on mappings and rules"""
    content_values = {}
    client_info = client_data.iloc[0]
    
    # Process field mappings
    for field, mapping in st.session_state.field_mappings.items():
        mapping_type = mapping.get("type")
        
        if mapping_type == "Simple Column Mapping":
            column = mapping.get("column")
            if column in client_data.columns:
                value = client_info[column]
                content_values[field] = format_handler.format_value(
                    value, column_name=column
                )
            else:
                content_values[field] = ""
        
        elif mapping_type == "Custom Text":
            text = mapping.get("text", "")
            # Replace {Column} placeholders
            for col in client_data.columns:
                placeholder = f"{{{col}}}"
                if placeholder in text:
                    col_value = client_info[col]
                    replacement = format_handler.format_value(
                        col_value, column_name=col
                    )
                    text = text.replace(placeholder, replacement)
            
            # Handle special placeholders
            if "{Date}" in text:
                text = text.replace("{Date}", datetime.now().strftime("%d %B %Y"))
            
            content_values[field] = text
        
        elif mapping_type == "Multi-Row Aggregation":
            agg_type = mapping.get("aggregation", "count")
            agg_column = mapping.get("agg_column", "")
            format_str = mapping.get("format", "{value}")
            
            # Apply filters if defined
            filtered_data = client_data
            if mapping.get("filter_column") and mapping.get("filter_operator"):
                filtered_data = apply_filter(
                    filtered_data, 
                    mapping["filter_column"], 
                    mapping["filter_operator"], 
                    mapping.get("filter_value", "")
                )
            
            # Perform aggregation
            if agg_type == "count":
                result = len(filtered_data)
            elif agg_type == "list" and agg_column in filtered_data.columns:
                values = filtered_data[agg_column].dropna().tolist()
                values = [str(val).strip() for val in values if str(val).strip()]
                
                list_style = mapping.get("list_style", "bullet")
                list_prefix = mapping.get("list_prefix", "• ")
                list_separator = mapping.get("list_separator", "\n")
                
                if list_separator == "\\n":
                    list_separator = "\n"
                
                if list_style == "bullet":
                    result = list_separator.join([f"{list_prefix}{item}" for item in values])
                elif list_style == "numbered":
                    result = list_separator.join([f"{i+1}. {item}" for i, item in enumerate(values)])
                elif list_style == "comma":
                    if len(values) > 1:
                        result = ", ".join(values[:-1]) + " and " + values[-1]
                    else:
                        result = values[0] if values else ""
                else:
                    result = list_separator.join(values)
            elif agg_column in filtered_data.columns:
                numeric_data = pd.to_numeric(filtered_data[agg_column], errors='coerce').dropna()
                if agg_type == "sum":
                    result = numeric_data.sum()
                elif agg_type == "average":
                    result = numeric_data.mean()
                elif agg_type == "min":
                    result = numeric_data.min()
                elif agg_type == "max":
                    result = numeric_data.max()
                else:
                    result = 0
            else:
                result = 0
            
            if isinstance(result, (int, float)):
                formatted_result = format_handler.format_value(
                    result, column_name=agg_column if agg_type != "count" else None
                )
            else:
                formatted_result = str(result)
            
            content_values[field] = format_str.replace("{value}", formatted_result)
    
    # Process conditional rules
    for field, rule in st.session_state.conditional_rules.items():
        result_text = evaluate_rule(rule, client_id, client_data, client_info)
        if result_text is not None and result_text != "{KEEP_ORIGINAL}":
            content_values[field] = result_text
    
    return content_values

def evaluate_rule(rule, client_id, client_data, client_info):
    """Evaluate a conditional rule"""
    rule_type = rule.get("type", "standard")
    
    # Check conditions
    condition_met = False
    
    if "conditions" in rule and isinstance(rule["conditions"], list):
        if rule_type == "standard":
            condition_met = evaluate_multiple_conditions(rule["conditions"], client_info)
        elif rule_type == "multi_row":
            aggregation = rule.get("aggregation", "any")
            if aggregation == "any":
                condition_met = any(
                    evaluate_multiple_conditions(rule["conditions"], row)
                    for _, row in client_data.iterrows()
                )
            elif aggregation == "all":
                if not client_data.empty:
                    condition_met = all(
                        evaluate_multiple_conditions(rule["conditions"], row)
                        for _, row in client_data.iterrows()
                    )
    else:
        # Legacy single condition
        column = rule.get("column", "")
        operator = rule.get("operator", "")
        value = rule.get("value", "")
        
        if column in client_info:
            condition_met = evaluate_condition(client_info[column], operator, value)
    
    # Return appropriate text
    if condition_met:
        if rule.get("remove_if_true", False):
            return None
        elif rule.get("keep_if_true", False):
            return "{KEEP_ORIGINAL}"
        else:
            return rule.get("then_text", "")
    else:
        if rule.get("remove_if_false", False):
            return None
        elif rule.get("keep_if_false", False):
            return "{KEEP_ORIGINAL}"
        else:
            return rule.get("else_text", "")

def create_word_table(table_def, client_data, format_handler):
    """Create a Word table based on table definition"""
    columns = table_def.get("columns", [])
    if not columns or client_data.empty:
        return ""
    
    # Get filtered data
    filtered_data = client_data.copy()
    filter_logic = table_def.get("filter_logic", "client_only")
    
    if filter_logic == "table_filter" and table_def.get("filter_column"):
        filtered_data = apply_filter(
            filtered_data,
            table_def["filter_column"],
            table_def.get("filter_operator", "equals"),
            table_def.get("filter_value", "")
        )
    elif filter_logic == "global_filters":
        for filter_dict in st.session_state.client_filters:
            if filter_dict.get("enabled", True):
                filtered_data = apply_filter(
                    filtered_data,
                    filter_dict["column"],
                    filter_dict["operator"],
                    filter_dict["value"]
                )
    
    # Build table HTML
    html = "<table border='1' style='border-collapse: collapse; width: 100%;'>"
    
    # Add title if requested
    if table_def.get("include_title", False):
        title_text = table_def.get("title_text", "")
        html += f"<tr><td colspan='{len(columns)}' style='text-align: center; font-weight: bold; background-color: #f0f0f0;'>{title_text}</td></tr>"
    
    # Add header if requested
    if table_def.get("include_header", True):
        html += "<tr style='background-color: #d3d3d3; font-weight: bold;'>"
        for col in columns:
            header = table_def.get("custom_headers", {}).get(col, col)
            html += f"<th style='padding: 8px; border: 1px solid #ccc;'>{header}</th>"
        html += "</tr>"
    
    # Add data rows
    for _, row in filtered_data.iterrows():
        html += "<tr>"
        for col in columns:
            value = row.get(col, "")
            formatted_value = format_handler.format_value(value, column_name=col)
            html += f"<td style='padding: 8px; border: 1px solid #ccc;'>{formatted_value}</td>"
        html += "</tr>"
    
    # Add totals row if requested
    if table_def.get("include_totals", False):
        totals_columns = table_def.get("totals_columns", [])
        if totals_columns:
            html += "<tr style='background-color: #f0f0f0; font-weight: bold;'>"
            html += f"<td style='padding: 8px; border: 1px solid #ccc;'>{table_def.get('totals_label', 'Totals')}</td>"
            
            for i, col in enumerate(columns[1:], 1):  # Skip first column (used for label)
                if col in totals_columns:
                    try:
                        total = pd.to_numeric(filtered_data[col], errors='coerce').sum()
                        formatted_total = format_handler.format_value(total, column_name=col)
                        html += f"<td style='padding: 8px; border: 1px solid #ccc;'>{formatted_total}</td>"
                    except:
                        html += "<td style='padding: 8px; border: 1px solid #ccc;'></td>"
                else:
                    html += "<td style='padding: 8px; border: 1px solid #ccc;'></td>"
            html += "</tr>"
    
    html += "</table>"
    return html

# --- Main Application ---

def main():
    initialize_session_state()
    
    st.title("📄 Advanced Document Assembly System")
    st.markdown("Professional document generation with advanced Excel integration, conditional logic, and table design.")
    
    # Sidebar Configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # File uploads
        uploaded_data = st.file_uploader("1. Upload Data Source", type=['xlsx', 'csv'])
        uploaded_template = st.file_uploader("2. Upload Word Template", type=['docx'])
        
        # Process uploaded data
        if uploaded_data and (not st.session_state.uploaded_data_name or uploaded_data.name != st.session_state.uploaded_data_name):
            with st.spinner("Loading data..."):
                if uploaded_data.name.endswith('.csv'):
                    st.session_state.df = pd.read_csv(uploaded_data)
                    st.session_state.excel_sheets = ["CSV"]
                    st.session_state.current_sheet = "CSV"
                    st.session_state.cell_formats = {}
                else:
                    file_content = uploaded_data.getvalue()
                    df, cell_formats, sheet_names = load_excel_with_formats(file_content)
                    if df is not None:
                        st.session_state.df = df
                        st.session_state.cell_formats = cell_formats
                        st.session_state.excel_sheets = sheet_names
                        st.session_state.current_sheet = sheet_names[0] if sheet_names else None
                
                st.session_state.uploaded_data_name = uploaded_data.name
                st.success("Data loaded successfully!")
        
        # Sheet selection for Excel
        if st.session_state.excel_sheets and len(st.session_state.excel_sheets) > 1:
            selected_sheet = st.selectbox("Select Sheet", st.session_state.excel_sheets)
            if selected_sheet != st.session_state.current_sheet:
                st.session_state.current_sheet = selected_sheet
                # Reload data for new sheet
                if uploaded_data and not uploaded_data.name.endswith('.csv'):
                    file_content = uploaded_data.getvalue()
                    df, cell_formats, _ = load_excel_with_formats(file_content, selected_sheet)
                    if df is not None:
                        st.session_state.df = df
                        st.session_state.cell_formats = cell_formats
        
        # Process uploaded template
        if uploaded_template and (not st.session_state.uploaded_template_name or uploaded_template.name != st.session_state.uploaded_template_name):
            with st.spinner("Analyzing template..."):
                st.session_state.detected_fields = detect_template_fields(uploaded_template)
                st.session_state.uploaded_template_name = uploaded_template.name
                
                if st.session_state.detected_fields:
                    st.success(f"Template analyzed! Found {len(st.session_state.detected_fields)} fields.")
                else:
                    st.warning("No template fields detected. Check the debug info in the Field Mapping tab.")
        
        # Debug info toggle
        if st.checkbox("Show Template Analysis Debug"):
            if st.session_state.template_analysis_debug:
                st.subheader("Template Analysis Debug")
                for debug_line in st.session_state.template_analysis_debug:
                    st.text(debug_line)
            else:
                st.text("No debug information available")
        
        # Client selection
        if st.session_state.df is not None:
            grouping_column = st.selectbox("3. Client Identifier Column", st.session_state.df.columns)
            
            # Client filters
            if st.session_state.client_filters:
                st.write("**Active Filters:**")
                for f in st.session_state.client_filters:
                    if f.get("enabled", True):
                        st.write(f"• {f['column']} {f['operator']} {f['value']}")
            
            # Apply filters to get available clients
            filtered_df = st.session_state.df.copy()
            for filter_dict in st.session_state.client_filters:
                if filter_dict.get("enabled", True):
                    filtered_df = apply_filter(
                        filtered_df,
                        filter_dict["column"],
                        filter_dict["operator"],
                        filter_dict["value"]
                    )
            
            available_clients = filtered_df[grouping_column].unique()
            
            process_option = st.radio("4. Processing Mode", 
                ("Single client (preview)", "All filtered clients"))
            
            if process_option.startswith("Single"):
                st.session_state.preview_client = st.selectbox("Select Client", available_clients)
        
        st.markdown("---")
        
        # Configuration save/load
        st.subheader("Save/Load Configuration")
        
        def save_config():
            return json.dumps({
                "field_mappings": st.session_state.field_mappings,
                "conditional_rules": st.session_state.conditional_rules,
                "client_filters": st.session_state.client_filters,
                "table_definitions": st.session_state.table_definitions,
                "on_the_fly_placeholders": st.session_state.on_the_fly_placeholders
            }, indent=2)
        
        st.download_button(
            "💾 Save Configuration",
            data=save_config(),
            file_name=f"doc_config_{datetime.now().strftime('%Y%m%d')}.json",
            mime="application/json"
        )
        
        uploaded_config = st.file_uploader("Load Configuration", type=['json'])
        if uploaded_config:
            try:
                config = json.load(uploaded_config)
                st.session_state.field_mappings = config.get("field_mappings", {})
                st.session_state.conditional_rules = config.get("conditional_rules", {})
                st.session_state.client_filters = config.get("client_filters", [])
                st.session_state.table_definitions = config.get("table_definitions", [])
                st.session_state.on_the_fly_placeholders = config.get("on_the_fly_placeholders", {})
                st.success("Configuration loaded!")
                st.rerun()
            except Exception as e:
                st.error(f"Error loading config: {e}")
    
    # Main content area
    if not uploaded_data or not uploaded_template or st.session_state.df is None:
        st.info("👆 Please upload both a data source and Word template to begin.")
        return
    
    # Tab navigation - Added new tab for on-the-fly placeholders
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "📋 Data Preview", 
        "🔗 Field Mapping", 
        "⚡ Conditional Rules", 
        "🗂️ Client Filters",
        "📊 Table Designer",
        "🏷️ On-the-Fly Placeholders",  # New tab
        "👀 Preview & Generate"
    ])
    
    with tab1:
        st.header("Data Preview")
        
        col1, col2 = st.columns([3, 1])
        with col1:
            st.dataframe(st.session_state.df.head(100), use_container_width=True)
        
        with col2:
            st.metric("Total Rows", len(st.session_state.df))
            st.metric("Columns", len(st.session_state.df.columns))
            if st.session_state.excel_sheets:
                st.metric("Current Sheet", st.session_state.current_sheet)
            
            st.subheader("Column Summary")
            for col in st.session_state.df.columns:
                non_null = st.session_state.df[col].notna().sum()
                st.write(f"**{col}**: {non_null}/{len(st.session_state.df)} values")
    
    with tab2:
        st.header("Field Mapping")
        
        # Template analysis results
        if st.session_state.template_analysis_debug:
            with st.expander("🔍 Template Analysis Details"):
                for debug_line in st.session_state.template_analysis_debug:
                    st.text(debug_line)
        
        if not st.session_state.detected_fields:
            st.warning("⚠️ No template fields detected!")
            st.info("""
            **Troubleshooting Tips:**
            1. Check that your Word template contains placeholders
            2. Supported formats: `{field}`, `{{field}}`, `[field]`, `[[field]]`, `<<<field>>>`, etc.
            3. Enable 'Show Template Analysis Debug' in the sidebar for detailed information
            4. Make sure the template file is not corrupted
            """)
        else:
            st.success(f"✅ Detected {len(st.session_state.detected_fields)} template fields")
            st.info("**Fields found:** " + ", ".join(st.session_state.detected_fields))
            
            # Field mapping interface
            for field in st.session_state.detected_fields:
                with st.expander(f"Configure Field: {field}", expanded=field not in st.session_state.field_mappings):
                    mapping_type = st.selectbox(
                        "Mapping Type", 
                        ["Simple Column Mapping", "Custom Text", "Multi-Row Aggregation"],
                        key=f"type_{field}"
                    )
                    
                    if mapping_type == "Simple Column Mapping":
                        column = st.selectbox("Data Column", [""] + list(st.session_state.df.columns), key=f"col_{field}")
                        format_type = st.selectbox("Format", ["Auto", "Currency (£)", "Percentage (%)", "Number", "Date", "Text"], key=f"fmt_{field}")
                        
                        if st.button(f"Save {field}", key=f"save_{field}"):
                            if column:
                                st.session_state.field_mappings[field] = {
                                    "type": mapping_type,
                                    "column": column,
                                    "format_type": format_type
                                }
                                st.success("Mapping saved!")
                    
                    elif mapping_type == "Custom Text":
                        text = st.text_area("Custom Text (use {ColumnName} for data)", key=f"text_{field}")
                        if st.button(f"Save {field}", key=f"save_{field}"):
                            st.session_state.field_mappings[field] = {
                                "type": mapping_type,
                                "text": text
                            }
                            st.success("Mapping saved!")
                    
                    elif mapping_type == "Multi-Row Aggregation":
                        col1, col2 = st.columns(2)
                        with col1:
                            agg_type = st.selectbox("Aggregation", ["count", "sum", "average", "min", "max", "list"], key=f"agg_{field}")
                            agg_column = st.selectbox("Column", [""] + list(st.session_state.df.columns), key=f"aggcol_{field}")
                        
                        with col2:
                            format_str = st.text_input("Format (use {value})", value="{value}", key=f"format_{field}")
                            
                            if agg_type == "list":
                                list_style = st.selectbox("List Style", ["bullet", "numbered", "comma", "line-break"], key=f"liststyle_{field}")
                        
                        # Filtering options
                        st.subheader("Filtering (Optional)")
                        filter_col = st.selectbox("Filter Column", [""] + list(st.session_state.df.columns), key=f"filtercol_{field}")
                        if filter_col:
                            filter_op = st.selectbox("Operator", ["equals", "contains", "starts with", "ends with", "greater than", "less than"], key=f"filterop_{field}")
                            filter_val = st.text_input("Value", key=f"filterval_{field}")
                        
                        if st.button(f"Save {field}", key=f"save_{field}"):
                            mapping = {
                                "type": mapping_type,
                                "aggregation": agg_type,
                                "agg_column": agg_column,
                                "format": format_str
                            }
                            
                            if agg_type == "list":
                                mapping.update({
                                    "list_style": list_style,
                                    "list_prefix": "• " if list_style == "bullet" else "",
                                    "list_separator": "\n"
                                })
                            
                            if filter_col:
                                mapping.update({
                                    "filter_column": filter_col,
                                    "filter_operator": filter_op,
                                    "filter_value": filter_val
                                })
                            
                            st.session_state.field_mappings[field] = mapping
                            st.success("Mapping saved!")
            
            # Show current mappings
            if st.session_state.field_mappings:
                st.subheader("Current Mappings")
                for field, mapping in st.session_state.field_mappings.items():
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        if mapping["type"] == "Simple Column Mapping":
                            st.write(f"**{field}** → {mapping.get('column', 'Not set')}")
                        elif mapping["type"] == "Custom Text":
                            st.write(f"**{field}** → Custom: {mapping.get('text', '')[:50]}...")
                        else:
                            st.write(f"**{field}** → {mapping.get('aggregation', 'count')}({mapping.get('agg_column', '')})")
                    
                    with col2:
                        if st.button("🗑️", key=f"del_{field}"):
                            del st.session_state.field_mappings[field]
                            st.rerun()
    
    with tab3:
        st.header("Conditional Rules")
        st.info("Create rules to insert different text based on data conditions.")
        
        # Add new rule interface
        with st.expander("➕ Add New Conditional Rule", expanded=False):
            rule_field = st.selectbox("Field to apply rule to", st.session_state.detected_fields)
            rule_type = st.selectbox("Rule Type", ["standard", "multi_row"])
            
            # Conditions builder
            st.subheader("Conditions")
            if 'new_rule_conditions' not in st.session_state:
                st.session_state.new_rule_conditions = [{"column": "", "operator": "", "value": "", "logic": "AND"}]
            
            for i, condition in enumerate(st.session_state.new_rule_conditions):
                col1, col2, col3, col4, col5 = st.columns([2, 2, 2, 1, 1])
                
                with col1:
                    if i > 0:
                        logic = st.selectbox("Logic", ["AND", "OR"], key=f"logic_{i}", index=0 if condition.get("logic") == "AND" else 1)
                        st.session_state.new_rule_conditions[i]["logic"] = logic
                    
                    column = st.selectbox("Column", [""] + list(st.session_state.df.columns), key=f"cond_col_{i}")
                    st.session_state.new_rule_conditions[i]["column"] = column
                
                with col2:
                    operator = st.selectbox("Operator", [
                        "equals", "does not equal", "contains", "does not contain",
                        "starts with", "ends with", "greater than", "less than",
                        "is empty", "is not empty"
                    ], key=f"cond_op_{i}")
                    st.session_state.new_rule_conditions[i]["operator"] = operator
                
                with col3:
                    value = st.text_input("Value", key=f"cond_val_{i}")
                    st.session_state.new_rule_conditions[i]["value"] = value
                
                with col4:
                    if i == len(st.session_state.new_rule_conditions) - 1:
                        if st.button("➕", key=f"add_cond_{i}"):
                            st.session_state.new_rule_conditions.append({"column": "", "operator": "", "value": "", "logic": "AND"})
                            st.rerun()
                
                with col5:
                    if len(st.session_state.new_rule_conditions) > 1:
                        if st.button("❌", key=f"del_cond_{i}"):
                            st.session_state.new_rule_conditions.pop(i)
                            st.rerun()
            
            # Rule actions
            st.subheader("Actions")
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("**IF condition is TRUE:**")
                then_action = st.radio("Action", ["Insert text", "Keep original", "Remove"], key="then_action")
                if then_action == "Insert text":
                    then_text = st.text_area("Text to insert", key="then_text")
            
            with col2:
                st.write("**IF condition is FALSE:**")
                else_action = st.radio("Action", ["Insert text", "Keep original", "Remove"], key="else_action")
                if else_action == "Insert text":
                    else_text = st.text_area("Text to insert", key="else_text")
            
            if st.button("💾 Save Rule"):
                rule = {
                    "type": rule_type,
                    "conditions": [c for c in st.session_state.new_rule_conditions if c["column"]],
                    "then_text": then_text if then_action == "Insert text" else "",
                    "remove_if_true": then_action == "Remove",
                    "keep_if_true": then_action == "Keep original",
                    "else_text": else_text if else_action == "Insert text" else "",
                    "remove_if_false": else_action == "Remove",
                    "keep_if_false": else_action == "Keep original"
                }
                
                st.session_state.conditional_rules[rule_field] = rule
                st.session_state.new_rule_conditions = [{"column": "", "operator": "", "value": "", "logic": "AND"}]
                st.success("Rule saved!")
                st.rerun()
        
        # Show existing rules
        if st.session_state.conditional_rules:
            st.subheader("Current Rules")
            for field, rule in st.session_state.conditional_rules.items():
                with st.expander(f"Rule for: {field}"):
                    conditions_text = " ".join([
                        f"{c.get('logic', '')} {c['column']} {c['operator']} '{c['value']}'" 
                        for c in rule.get("conditions", [])
                    ]).strip()
                    
                    st.write(f"**Conditions:** {conditions_text}")
                    st.write(f"**Then:** {rule.get('then_text', 'N/A')}")
                    st.write(f"**Else:** {rule.get('else_text', 'N/A')}")
                    
                    if st.button(f"🗑️ Delete Rule", key=f"del_rule_{field}"):
                        del st.session_state.conditional_rules[field]
                        st.rerun()
    
    with tab4:
        st.header("Client Filters")
        st.info("Filter which clients will receive documents.")
        
        # Add new filter
        with st.expander("➕ Add New Filter", expanded=False):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                filter_column = st.selectbox("Column", st.session_state.df.columns, key="new_filter_col")
            
            with col2:
                filter_operator = st.selectbox("Operator", [
                    "equals", "contains", "starts with", "ends with",
                    "greater than", "less than", "is empty", "is not empty"
                ], key="new_filter_op")
            
            with col3:
                filter_value = st.text_input("Value", key="new_filter_val")
            
            if st.button("Add Filter"):
                new_filter = {
                    "column": filter_column,
                    "operator": filter_operator,
                    "value": filter_value,
                    "enabled": True
                }
                st.session_state.client_filters.append(new_filter)
                st.rerun()
        
        # Show current filters
        if st.session_state.client_filters:
            st.subheader("Active Filters")
            
            # Calculate filter results
            original_count = len(st.session_state.df[grouping_column].unique()) if 'grouping_column' in locals() else 0
            filtered_df = st.session_state.df.copy()
            
            for filter_dict in st.session_state.client_filters:
                if filter_dict.get("enabled", True):
                    filtered_df = apply_filter(
                        filtered_df,
                        filter_dict["column"],
                        filter_dict["operator"],
                        filter_dict["value"]
                    )
            
            filtered_count = len(filtered_df[grouping_column].unique()) if 'grouping_column' in locals() else 0
            
            st.metric("Clients After Filtering", f"{filtered_count} / {original_count}")
            
            for i, f in enumerate(st.session_state.client_filters):
                col1, col2, col3 = st.columns([1, 4, 1])
                
                with col1:
                    enabled = st.checkbox("", value=f.get("enabled", True), key=f"filter_enabled_{i}")
                    st.session_state.client_filters[i]["enabled"] = enabled
                
                with col2:
                    st.write(f"{f['column']} {f['operator']} '{f['value']}'")
                
                with col3:
                    if st.button("🗑️", key=f"del_filter_{i}"):
                        st.session_state.client_filters.pop(i)
                        st.rerun()
    
    with tab5:
        st.header("Table Designer")
        st.info("Design tables that will be inserted into your documents.")
        
        # Add new table
        with st.expander("➕ Add New Table Definition", expanded=False):
            table_name = st.text_input("Table Name")
            table_placeholder = st.text_input("Placeholder (e.g., {TABLE:ClientData})")
            
            col1, col2 = st.columns(2)
            with col1:
                include_header = st.checkbox("Include Header Row", value=True)
                include_title = st.checkbox("Include Title Row")
                if include_title:
                    title_text = st.text_input("Title Text")
            
            with col2:
                include_totals = st.checkbox("Include Totals Row")
                if include_totals:
                    totals_label = st.text_input("Totals Label", value="Totals")
            
            # Column selection
            selected_columns = st.multiselect("Select Columns", st.session_state.df.columns)
            
            # Custom headers
            custom_headers = {}
            if selected_columns:
                st.subheader("Custom Headers (Optional)")
                for col in selected_columns:
                    custom_header = st.text_input(f"Header for {col}", value=col, key=f"header_{col}")
                    if custom_header != col:
                        custom_headers[col] = custom_header
            
            # Totals columns
            totals_columns = []
            if include_totals and selected_columns:
                totals_columns = st.multiselect("Columns to Total", selected_columns)
            
            if st.button("💾 Save Table Definition"):
                table_def = {
                    "name": table_name,
                    "placeholder": table_placeholder,
                    "columns": selected_columns,
                    "include_header": include_header,
                    "include_title": include_title,
                    "title_text": title_text if include_title else "",
                    "include_totals": include_totals,
                    "totals_label": totals_label if include_totals else "Totals",
                    "totals_columns": totals_columns,
                    "custom_headers": custom_headers,
                    "filter_logic": "client_only"
                }
                
                st.session_state.table_definitions.append(table_def)
                st.success("Table definition saved!")
                st.rerun()
        
        # Show existing tables
        if st.session_state.table_definitions:
            st.subheader("Defined Tables")
            for i, table_def in enumerate(st.session_state.table_definitions):
                with st.expander(f"Table: {table_def['name']}"):
                    st.write(f"**Placeholder:** {table_def['placeholder']}")
                    st.write(f"**Columns:** {', '.join(table_def['columns'])}")
                    st.write(f"**Features:** Header: {table_def['include_header']}, Title: {table_def['include_title']}, Totals: {table_def['include_totals']}")
                    
                    # Preview table
                    if st.session_state.preview_client and 'grouping_column' in locals():
                        client_data = st.session_state.df[st.session_state.df[grouping_column] == st.session_state.preview_client]
                        if not client_data.empty:
                            table_html = create_word_table(table_def, client_data, st.session_state.format_handler)
                            st.markdown("**Preview:**")
                            st.markdown(table_html, unsafe_allow_html=True)
                    
                    if st.button(f"🗑️ Delete Table", key=f"del_table_{i}"):
                        st.session_state.table_definitions.pop(i)
                        st.rerun()
    
    # NEW TAB: On-the-Fly Placeholders
    with tab6:
        st.header("🏷️ On-the-Fly Placeholders")
        st.info("Create custom placeholders that will be dynamically inserted into your template before processing. This allows you to add new placeholders without manually editing the Word template.")
        
        # Add new placeholder
        with st.expander("➕ Add New On-the-Fly Placeholder", expanded=False):
            col1, col2 = st.columns(2)
            
            with col1:
                placeholder_name = st.text_input("Placeholder Name", help="A unique name for this placeholder")
                replacement_text = st.text_input("Text to Replace in Template", help="The text in your template that will be replaced with the placeholder")
            
            with col2:
                placeholder_format = st.selectbox("Placeholder Format", [
                    "{{placeholder_name}}", 
                    "{placeholder_name}", 
                    "[placeholder_name]", 
                    "[[placeholder_name]]",
                    "<<<placeholder_name>>>",
                    "<<placeholder_name>>",
                    "<placeholder_name>"
                ], help="Choose the format for the placeholder that will be inserted")
                
                # Show preview of actual placeholder
                if placeholder_name:
                    actual_placeholder = placeholder_format.replace("placeholder_name", placeholder_name)
                    st.write(f"**Actual placeholder:** `{actual_placeholder}`")
            
            st.markdown("---")
            st.subheader("Placeholder Mapping")
            st.info("Configure how this placeholder should be filled with data (same as regular field mapping)")
            
            # Mapping configuration (similar to field mapping)
            mapping_type = st.selectbox("Mapping Type", 
                ["Simple Column Mapping", "Custom Text", "Multi-Row Aggregation"],
                key="otf_mapping_type"
            )
            
            mapping_config = {}
            
            if mapping_type == "Simple Column Mapping":
                column = st.selectbox("Data Column", [""] + list(st.session_state.df.columns), key="otf_col")
                format_type = st.selectbox("Format", ["Auto", "Currency (£)", "Percentage (%)", "Number", "Date", "Text"], key="otf_fmt")
                mapping_config = {
                    "type": mapping_type,
                    "column": column,
                    "format_type": format_type
                }
            
            elif mapping_type == "Custom Text":
                text = st.text_area("Custom Text (use {ColumnName} for data)", key="otf_text")
                mapping_config = {
                    "type": mapping_type,
                    "text": text
                }
            
            elif mapping_type == "Multi-Row Aggregation":
                col1_inner, col2_inner = st.columns(2)
                with col1_inner:
                    agg_type = st.selectbox("Aggregation", ["count", "sum", "average", "min", "max", "list"], key="otf_agg")
                    agg_column = st.selectbox("Column", [""] + list(st.session_state.df.columns), key="otf_aggcol")
                
                with col2_inner:
                    format_str = st.text_input("Format (use {value})", value="{value}", key="otf_format")
                    
                    if agg_type == "list":
                        list_style = st.selectbox("List Style", ["bullet", "numbered", "comma", "line-break"], key="otf_liststyle")
                
                mapping_config = {
                    "type": mapping_type,
                    "aggregation": agg_type,
                    "agg_column": agg_column,
                    "format": format_str
                }
                
                if agg_type == "list":
                    mapping_config.update({
                        "list_style": list_style,
                        "list_prefix": "• " if list_style == "bullet" else "",
                        "list_separator": "\n"
                    })
            
            if st.button("💾 Save On-the-Fly Placeholder"):
                if placeholder_name and replacement_text:
                    actual_placeholder = placeholder_format.replace("placeholder_name", placeholder_name)
                    
                    st.session_state.on_the_fly_placeholders[placeholder_name] = {
                        "replacement_text": replacement_text,
                        "placeholder_format": actual_placeholder,
                        "mapping": mapping_config
                    }
                    
                    # Also add to field mappings so it gets processed
                    st.session_state.field_mappings[placeholder_name] = mapping_config
                    
                    st.success("On-the-fly placeholder saved!")
                    st.rerun()
                else:
                    st.error("Please provide both placeholder name and replacement text")
        
        # Show existing on-the-fly placeholders
        if st.session_state.on_the_fly_placeholders:
            st.subheader("Current On-the-Fly Placeholders")
            
            for placeholder_name, placeholder_info in st.session_state.on_the_fly_placeholders.items():
                with st.expander(f"Placeholder: {placeholder_name}"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write(f"**Replaces text:** `{placeholder_info['replacement_text']}`")
                        st.write(f"**Placeholder format:** `{placeholder_info['placeholder_format']}`")
                    
                    with col2:
                        mapping = placeholder_info.get('mapping', {})
                        if mapping.get('type') == 'Simple Column Mapping':
                            st.write(f"**Maps to:** {mapping.get('column', 'Not set')}")
                        elif mapping.get('type') == 'Custom Text':
                            st.write(f"**Custom text:** {mapping.get('text', '')[:50]}...")
                        elif mapping.get('type') == 'Multi-Row Aggregation':
                            st.write(f"**Aggregation:** {mapping.get('aggregation', 'count')}({mapping.get('agg_column', '')})")
                    
                    if st.button(f"🗑️ Delete", key=f"del_otf_{placeholder_name}"):
                        # Remove from both dictionaries
                        del st.session_state.on_the_fly_placeholders[placeholder_name]
                        if placeholder_name in st.session_state.field_mappings:
                            del st.session_state.field_mappings[placeholder_name]
                        st.rerun()
            
            # Show preview of template modifications
            st.subheader("Template Modification Preview")
            st.info("This shows how your template text will be modified before processing:")
            
            for placeholder_name, placeholder_info in st.session_state.on_the_fly_placeholders.items():
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Original:** `{placeholder_info['replacement_text']}`")
                with col2:
                    st.write(f"**Becomes:** `{placeholder_info['placeholder_format']}`")
    
    with tab7:
        st.header("Preview & Generate Documents")
        
        # Preview section
        if st.session_state.preview_client and 'grouping_column' in locals():
            st.subheader(f"Preview for Client: {st.session_state.preview_client}")
            
            # Get client data
            client_data = st.session_state.df[st.session_state.df[grouping_column] == st.session_state.preview_client]
            if not client_data.empty:
                # Generate content values
                content_values = generate_content_values(
                    st.session_state.preview_client, 
                    client_data, 
                    st.session_state.format_handler
                )
                
                # Show field values
                with st.expander("Field Values", expanded=True):
                    for field, value in content_values.items():
                        st.write(f"**{field}:** {value}")
                
                # Show on-the-fly placeholders if any
                if st.session_state.on_the_fly_placeholders:
                    with st.expander("On-the-Fly Placeholders", expanded=True):
                        for placeholder_name, placeholder_info in st.session_state.on_the_fly_placeholders.items():
                            value = content_values.get(placeholder_name, "N/A")
                            st.write(f"**{placeholder_name}** (`{placeholder_info['replacement_text']}` → `{placeholder_info['placeholder_format']}`): {value}")
                
                # Show tables
                if st.session_state.table_definitions:
                    with st.expander("Tables", expanded=True):
                        for table_def in st.session_state.table_definitions:
                            st.write(f"**{table_def['name']}** ({table_def['placeholder']})")
                            table_html = create_word_table(table_def, client_data, st.session_state.format_handler)
                            st.markdown(table_html, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Generation section
        st.subheader("Generate Documents")
        
        if not st.session_state.field_mappings and not st.session_state.on_the_fly_placeholders:
            st.warning("No field mappings or on-the-fly placeholders configured. Please set up at least one type of mapping first.")
        else:
            col1, col2 = st.columns(2)
            
            with col1:
                filename_template = st.text_input("Filename Template", value="{client_id}_report_{date}")
            
            with col2:
                output_format = st.selectbox("Output Format", ["Word (.docx)", "PDF", "Both"])
            
            if st.button("🚀 Generate Documents", type="primary"):
                with st.spinner("Generating documents..."):
                    try:
                        # Determine clients to process
                        if process_option.startswith("Single"):
                            clients_to_process = [st.session_state.preview_client]
                        else:
                            # Apply filters
                            filtered_df = st.session_state.df.copy()
                            for filter_dict in st.session_state.client_filters:
                                if filter_dict.get("enabled", True):
                                    filtered_df = apply_filter(
                                        filtered_df,
                                        filter_dict["column"],
                                        filter_dict["operator"],
                                        filter_dict["value"]
                                    )
                            clients_to_process = filtered_df[grouping_column].unique()
                        
                        # Progress tracking
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        generated_files = {}
                        failed_clients = []
                        
                        for i, client_id in enumerate(clients_to_process):
                            status_text.text(f"Processing client {i+1}/{len(clients_to_process)}: {client_id}")
                            
                            try:
                                # Get client data
                                client_data = st.session_state.df[st.session_state.df[grouping_column] == client_id]
                                
                                # Generate content values
                                content_values = generate_content_values(
                                    client_id, 
                                    client_data, 
                                    st.session_state.format_handler
                                )
                                
                                # Add table data
                                for table_def in st.session_state.table_definitions:
                                    placeholder = table_def['placeholder']
                                    table_html = create_word_table(table_def, client_data, st.session_state.format_handler)
                                    content_values[placeholder] = table_html
                                
                                # Start with original template
                                uploaded_template.seek(0)  # Reset file position
                                template_buffer = io.BytesIO(uploaded_template.read())
                                
                                # STEP 1: Apply on-the-fly placeholders if any
                                if st.session_state.on_the_fly_placeholders:
                                    # Load the document to modify text
                                    temp_doc = docx.Document(template_buffer)
                                    
                                    # Replace text in paragraphs
                                    for paragraph in temp_doc.paragraphs:
                                        for placeholder_name, placeholder_info in st.session_state.on_the_fly_placeholders.items():
                                            replacement_text = placeholder_info['replacement_text']
                                            placeholder_format = placeholder_info['placeholder_format']
                                            if replacement_text in paragraph.text:
                                                paragraph.text = paragraph.text.replace(replacement_text, placeholder_format)
                                    
                                    # Replace text in table cells
                                    for table in temp_doc.tables:
                                        for row in table.rows:
                                            for cell in row.cells:
                                                for paragraph in cell.paragraphs:
                                                    for placeholder_name, placeholder_info in st.session_state.on_the_fly_placeholders.items():
                                                        replacement_text = placeholder_info['replacement_text']
                                                        placeholder_format = placeholder_info['placeholder_format']
                                                        if replacement_text in paragraph.text:
                                                            paragraph.text = paragraph.text.replace(replacement_text, placeholder_format)
                                    
                                    # Save modified template to buffer
                                    modified_template_buffer = io.BytesIO()
                                    temp_doc.save(modified_template_buffer)
                                    modified_template_buffer.seek(0)
                                    template_buffer = modified_template_buffer
                                else:
                                    template_buffer.seek(0)
                                
                                # STEP 2: Convert all non-Jinja placeholders to Jinja format
                                template_buffer = convert_placeholders_to_jinja(template_buffer, st.session_state.detected_fields)
                                
                                # STEP 3: Generate document using DocxTemplate (now all placeholders are Jinja-compatible)
                                doc = DocxTemplate(template_buffer)
                                
                                # Add client_data for table loops
                                content_values['client_data'] = client_data.to_dict('records')
                                content_values['report_date'] = datetime.now().strftime("%d %B %Y")
                                
                                # Render the document
                                doc.render(content_values)
                                
                                # Save to buffer
                                doc_buffer = io.BytesIO()
                                doc.save(doc_buffer)
                                doc_bytes = doc_buffer.getvalue()
                                
                                # Generate filename
                                filename = filename_template.format(
                                    client_id=client_id,
                                    date=datetime.now().strftime("%Y%m%d")
                                )
                                filename = re.sub(r'[\\/*?:"<>|]', "_", filename)
                                
                                generated_files[f"{filename}.docx"] = doc_bytes
                                
                            except Exception as e:
                                failed_clients.append(f"{client_id}: {str(e)}")
                            
                            progress_bar.progress((i + 1) / len(clients_to_process))
                        
                        status_text.text("Generation complete!")
                        
                        # Results
                        if generated_files:
                            st.success(f"✅ Generated {len(generated_files)} documents successfully!")
                            
                            if len(generated_files) == 1:
                                # Single file download
                                filename, file_data = list(generated_files.items())[0]
                                st.download_button(
                                    f"📥 Download {filename}",
                                    data=file_data,
                                    file_name=filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            else:
                                # Multiple files - create zip
                                zip_buffer = io.BytesIO()
                                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                                    for filename, file_data in generated_files.items():
                                        zf.writestr(filename, file_data)
                                
                                st.download_button(
                                    "📥 Download All Documents (ZIP)",
                                    data=zip_buffer.getvalue(),
                                    file_name=f"documents_{datetime.now().strftime('%Y%m%d')}.zip",
                                    mime="application/zip"
                                )
                            
                            # Statistics
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("Total Processed", len(clients_to_process))
                            with col2:
                                st.metric("Successful", len(generated_files))
                            with col3:
                                st.metric("Failed", len(failed_clients))
                            
                            # Show failures if any
                            if failed_clients:
                                with st.expander("Failed Clients"):
                                    for failure in failed_clients:
                                        st.error(failure)
                        
                        else:
                            st.error("No documents were generated successfully.")
                            if failed_clients:
                                for failure in failed_clients:
                                    st.error(failure)
                    
                    except Exception as e:
                        st.error(f"Generation failed: {str(e)}")

if __name__ == "__main__":
    main()